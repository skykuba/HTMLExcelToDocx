from html2docx import html2docx
import os
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

TEMPLATE_PATH = 'template.html'
OUTPUT_FOLDER = 'words/'
TEMP_HTML_PATH = 'temp.html'
DATABASE_PATH = 'dane/testDane.xlsx'

REPLACE_IN_HTML = ['@title', '@opis', '@nazwa_sql', '@nazwa_dex', '@struktury',
                   '@klucz']  # zakladamy ze tytul jest unikalny i jest on nazwa worda
HEADERS_IN_WORD = ['Opis', 'Nazwa SQL', 'Nazwa Dex', 'Opis Struktury', 'Klucz']
HEADERS_IN_EXCEL = ['Opis', 'NazwaSQL', 'NazwaDex', 'OpisStrukutury', 'Klucz']
FONT_NAME = 'Arial'

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

database_path = pd.read_excel(DATABASE_PATH, sheet_name='Arkusz1')

output_folder = OUTPUT_FOLDER
if not os.path.exists(output_folder):
    os.makedirs(output_folder)


def is_valid(value):
    return value not in ["<p><br></p>", "-", "<p></p>", " ", ""]


def process_content(content):
    return f'<p>{content}</p>' if is_valid(content) else '<p></p>'


# funkcja generuje tymczasowy plik html na podstawie szablonu html i argumentów
def to_html(title, arguments, html_temp_path=TEMP_HTML_PATH):
    with open(TEMPLATE_PATH, 'r') as file:
        template_html = file.read()

    title_html = '<h1>' + title + '</h1>'
    content_to_replace = [title_html] + [process_content(arg) for arg in arguments]

    if len(content_to_replace) != len(REPLACE_IN_HTML):
        raise ValueError('Niepoprawana liczba argumentów do zastąpienia w szablonie HTML')

    temp_html = template_html
    for i in range(len(REPLACE_IN_HTML)):
        temp_html = temp_html.replace(REPLACE_IN_HTML[i], content_to_replace[i], 1)

    with open(html_temp_path, 'w') as file:
        file.write(temp_html)


# funkcja konwertuje plik html na plik docx
def from_html_to_docx(html_temp_path, output_path, title):
    with open(html_temp_path) as fp:
        html = fp.read()
    buf = html2docx(html, title=title)
    output_path = output_path + title + ".docx"
    with open(output_path, "wb") as fp:
        fp.write(buf.getvalue())


# funkcja dodaje cieniowanie do paragrafu
def add_paragraph_shading(paragraph, color="D9D9D9"):
    shading_element = OxmlElement('w:shd')
    shading_element.set(qn('w:fill'), color)
    parag_pr = paragraph._element.get_or_add_pPr()
    parag_pr.append(shading_element)


# funkcja zwiększa czcionkę w dokumencie docx
def increase_font(docx_path, title, font_name=FONT_NAME):
    search_texts = [title] + HEADERS_IN_WORD
    doc = Document(docx_path)
    collected_texts = {text: False for text in search_texts}

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for search_text in search_texts:
                if search_text in run.text and not collected_texts.get(search_text, False):
                    run.font.name = font_name
                    if run.text == title:
                        run.font.size = Pt(14)
                    elif search_text in HEADERS_IN_WORD:
                        run.font.size = Pt(12)
                    collected_texts[search_text] = True
                    run.font.color.rgb = RGBColor(000, 150, 143)
                    break
                else:
                    run.font.name = font_name
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(000, 000, 000)
    doc.save(docx_path)


i = 1

for index, row in database_path.iterrows():
    tytul = str(row['Tytul'])
    if pd.isna(row['Tytul']):
        print('Tytuł to NaN, przerywam przetwarzanie.')
        break
    forbidden_chars_pattern = r'[<>:"/\\|?*]+'
    sanitized_title = re.sub(forbidden_chars_pattern, '_', tytul)
    tytul = sanitized_title
    arguments = [row[header] for header in HEADERS_IN_EXCEL]
    to_html(tytul, arguments, TEMP_HTML_PATH)

    from_html_to_docx(TEMP_HTML_PATH, OUTPUT_FOLDER, tytul)
    increase_font(OUTPUT_FOLDER + tytul + ".docx", tytul)

    print(str(i) + ' udało się: ' + tytul)
    i += 1
print('Udało się wszystko')
